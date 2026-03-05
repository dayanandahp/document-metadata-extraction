"""
Text Extraction Module
Handles extracting text from .docx and .png files
Enhanced with advanced image preprocessing for better OCR accuracy
"""

import os
import numpy as np
from docx import Document
from PIL import Image
import pytesseract
import platform

try:
    import cv2
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False
    print(" opencv-python not installed. Using basic image preprocessing only.")

# Set Tesseract path for Windows
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    )


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a .docx file
    
    Args:
        file_path: Path to the .docx file
    
    Returns:
        Extracted text as string
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                full_text.append(para.text.strip())
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        return '\n'.join(full_text)
    
    except Exception as e:
        print(f"Error extracting text from docx {file_path}: {e}")
        return ""


def preprocess_image_cv2(file_path: str) -> list:
    """
    Advanced image preprocessing using OpenCV.
    Returns multiple preprocessed versions to try with OCR.
    """
    img = cv2.imread(file_path)
    if img is None:
        return []
    
    preprocessed = []
    
    # Version 1: Grayscale + resize + adaptive threshold
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # Resize to 2x for better OCR on small text
    gray_resized = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    # Denoise
    denoised = cv2.fastNlMeansDenoising(gray_resized, h=30)
    # Adaptive threshold
    binary_adaptive = cv2.adaptiveThreshold(
        denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 31, 2
    )
    preprocessed.append(binary_adaptive)
    
    # Version 2: Otsu's thresholding (good for bimodal images)
    _, binary_otsu = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    preprocessed.append(binary_otsu)
    
    # Version 3: Just grayscale + denoise (no thresholding)
    preprocessed.append(denoised)
    
    # Version 4: Sharpen + threshold
    kernel_sharpen = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
    sharpened = cv2.filter2D(gray_resized, -1, kernel_sharpen)
    _, binary_sharp = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    preprocessed.append(binary_sharp)
    
    return preprocessed


def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from an image file using OCR.
    Uses advanced preprocessing (OpenCV) and tries multiple Tesseract PSM modes
    to maximize extraction quality.
    
    Args:
        file_path: Path to the image file (.png, .jpg)
    
    Returns:
        Extracted text as string
    """
    try:
        best_text = ""
        
        # PSM modes to try:
        # 3 = Fully automatic page segmentation (default)
        # 4 = Assume single column of text
        # 6 = Assume a single uniform block of text
        # 1 = Automatic with OSD
        psm_modes = [6, 4, 3, 1]
        
        if HAS_CV2:
            # Use advanced OpenCV preprocessing
            preprocessed_images = preprocess_image_cv2(file_path)
            
            for img_array in preprocessed_images:
                # Convert numpy array to PIL Image for pytesseract
                pil_img = Image.fromarray(img_array)
                
                for psm in psm_modes:
                    try:
                        config = f'--psm {psm} --oem 3'
                        text = pytesseract.image_to_string(pil_img, config=config)
                        text = text.strip()
                        
                        # Pick the version that extracts the most text
                        if len(text) > len(best_text):
                            best_text = text
                    except Exception:
                        continue
        
        # Also try basic PIL approach (as fallback or if cv2 not available)
        image = Image.open(file_path)
        image_gray = image.convert('L')
        
        for psm in psm_modes:
            try:
                config = f'--psm {psm} --oem 3'
                text = pytesseract.image_to_string(image_gray, config=config)
                text = text.strip()
                if len(text) > len(best_text):
                    best_text = text
            except Exception:
                continue
        
        if best_text:
            print(f"   OCR extracted {len(best_text)} characters")
        else:
            print(f"   OCR extracted no text")
        
        return best_text
    
    except Exception as e:
        print(f"Error extracting text from image {file_path}: {e}")
        return ""


def extract_text(file_path: str) -> str:
    """
    Main function: Route to appropriate extractor based on file type
    
    Args:
        file_path: Path to the document
    
    Returns:
        Extracted text as string
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        print(f"   Extracting text from DOCX: {os.path.basename(file_path)}")
        return extract_text_from_docx(file_path)
    
    elif ext in ['.png', '.jpg', '.jpeg']:
        print(f"   Extracting text from Image (OCR): {os.path.basename(file_path)}")
        return extract_text_from_image(file_path)
    
    else:
        print(f"   Unsupported file type: {ext}")
        return ""


# ===== TEST THIS MODULE =====
if __name__ == "__main__":
    # Quick test
    import sys
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        text = extract_text(test_file)
        print("\n--- EXTRACTED TEXT ---")
        print(text[:1000])  # Print first 1000 chars
        print(f"\n--- Total length: {len(text)} characters ---")
    else:
        print("Usage: python -m src.text_extractor <file_path>")